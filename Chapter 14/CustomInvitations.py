#! Python3
# CustomInvitations.py - Reads a text document of guests and creates custom invitations in multiple word documents
# Use --> python3 CustomInvitations.py (name of guest list text document)

import docx, sys

def main():
    if len(sys.argv) != 2:
        print('Invalid number of arguments!')
        print('Use --> python3 CustomInvitations.py (name of guest list text document)')
        exit(0)

    createCustomInvitations(sys.argv[1])

def createCustomInvitations(fileName):
    # Create a list of guests from the text file
    guestListFile = open(fileName, 'r')
    guestList = guestListFile.readlines()
    # Open the format docx file
    doc = docx.Document('Format.docx')
    # Add invitations to docx file with 1 invitation per page
    for name in guestList:
        doc.add_paragraph('It would be the pleasure to have the company of', 'Cursive')
        doc.add_paragraph(name.strip(), 'Name')
        doc.add_paragraph('at 11010 Memorial Lane on the Evening of', 'Cursive')
        doc.add_paragraph('April 1st', 'Date')
        doc.add_paragraph("at 7 o'clock", 'Cursive')
        doc.paragraphs[len(doc.paragraphs) - 1].runs[len(doc.paragraphs[len(doc.paragraphs) - 1].runs) - 1].add_break(docx.enum.text.WD_BREAK.PAGE)
    
    # Save as new file
    doc.save('Invitations.docx')
    guestListFile.close()

if __name__ == '__main__':
    main()